#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Export experiment bundle: configs, eval JSON, aggregates, JSONL, CSV index.

Outputs (under experiments/):
  EXPERIMENT_DATA_BUNDLE.md   (--format md, default)
  EXPERIMENT_DATA_BUNDLE.json (--format json)
  EXPERIMENT_DATA_BUNDLE.docx (--format docx)
  EXPERIMENT_DATA_BUNDLE.pdf  (--format pdf)

  --format all  writes md + json + docx + pdf

Run from repo root:
  python experiments/export_experiment_data_bundle.py
  python experiments/export_experiment_data_bundle.py --format pdf
  pip install python-docx reportlab   # 可选；无 reportlab 时用 matplotlib 生成 PDF
"""
from __future__ import annotations

import argparse
import json
import os
import platform
from datetime import datetime, timezone
from io import StringIO
from pathlib import Path
from typing import Any, Dict, List, Sequence

import numpy as np
import yaml

ROOT = Path(__file__).resolve().parents[1]
RESULTS_EVAL = ROOT / "experiments" / "results_eval"
SPEC_PATH = ROOT / "experiments" / "aggregate_dual_spec.json"
OUT_MD = ROOT / "experiments" / "EXPERIMENT_DATA_BUNDLE.md"
OUT_JSON = ROOT / "experiments" / "EXPERIMENT_DATA_BUNDLE.json"
OUT_DOCX = ROOT / "experiments" / "EXPERIMENT_DATA_BUNDLE.docx"
OUT_PDF = ROOT / "experiments" / "EXPERIMENT_DATA_BUNDLE.pdf"

YAML_SNAPSHOTS = [
    "configs/envs/grid_3d_medium.yaml",
    "configs/envs/grid_3d_hard.yaml",
    "configs/algos/qmix_paper.yaml",
    "configs/algos/vdn.yaml",
    "configs/algos/qmix_no_potential.yaml",
    "configs/algos/mappo.yaml",
]

JSONL_INLINE = [
    "experiments/results/medium_qmix_vdn.jsonl",
    "experiments/results/hard_qmix_vdn.jsonl",
    "experiments/results/hard_qmix_only.jsonl",
    "experiments/results/ablation_medium.jsonl",
    "experiments/results/curve_aux.jsonl",
    "experiments/results/hard_curve_aux.jsonl",
    "experiments/results/single_lines.jsonl",
]


def _load_json(path: Path) -> Dict[str, Any]:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def _sanitize_eval_record(data: Dict[str, Any]) -> Dict[str, Any]:
    out = dict(data)
    ck = out.get("checkpoint")
    if isinstance(ck, str) and ck:
        out["checkpoint"] = Path(ck).name
    return out


def _resolve(p: str) -> Path:
    path = Path(p)
    return path if path.is_absolute() else (ROOT / path).resolve()


def _aggregate_row(
    name: str, paths: Sequence[Path], metrics: List[str]
) -> Dict[str, Any]:
    per_seed: Dict[str, List[float]] = {m: [] for m in metrics}
    seed_files: List[str] = []
    for path in paths:
        seed_files.append(str(path.relative_to(ROOT)))
        if not path.is_file():
            for m in metrics:
                per_seed[m].append(float("nan"))
            continue
        data = _load_json(path)
        for m in metrics:
            v = data.get(m)
            per_seed[m].append(float("nan") if v is None else float(v))

    summary: Dict[str, Any] = {"method": name, "seed_files": seed_files}
    for m in metrics:
        arr = np.array(per_seed[m], dtype=float)
        n = int(np.sum(~np.isnan(arr)))
        summary[m] = {
            "per_seed": [None if np.isnan(x) else float(x) for x in arr.tolist()],
            "mean": float(np.nanmean(arr)) if n else None,
            "std_sample": float(np.nanstd(arr, ddof=1)) if n > 1 else 0.0,
            "n": n,
        }
    return summary


def collect_bundle() -> Dict[str, Any]:
    bundle: Dict[str, Any] = {
        "bundle_format": "uav-qmix-3d-experiment-data",
        "generated_utc": datetime.now(timezone.utc).isoformat(),
        "repo_root": str(ROOT),
        "description_zh": (
            "汇总：环境/算法 YAML、每种子评估 JSON、双档五方法聚合、"
            "小体积训练 jsonl 与曲线 CSV 索引。"
        ),
    }

    bundle["config_snapshots_text"] = {}
    for rel in YAML_SNAPSHOTS:
        p = ROOT / rel
        if p.is_file():
            bundle["config_snapshots_text"][rel.replace("\\", "/")] = p.read_text(
                encoding="utf-8"
            )

    bundle["config_snapshots_yaml"] = {}
    for rel in YAML_SNAPSHOTS:
        p = ROOT / rel
        if p.is_file():
            bundle["config_snapshots_yaml"][rel.replace("\\", "/")] = yaml.safe_load(
                p.read_text(encoding="utf-8")
            )

    bundle["results_eval"] = {}
    if RESULTS_EVAL.is_dir():
        for path in sorted(RESULTS_EVAL.glob("*.json")):
            try:
                bundle["results_eval"][path.name] = _sanitize_eval_record(
                    _load_json(path)
                )
            except Exception as exc:  # noqa: BLE001
                bundle["results_eval"][path.name] = {"_error": str(exc)}

    metrics = ["coverage_mean", "success_mean", "steps_mean"]
    bundle["aggregated_dual_table"] = {"metrics": metrics, "tiers": []}

    if SPEC_PATH.is_file():
        spec = _load_json(SPEC_PATH)
        for tier in spec.get("tiers", []):
            title = str(tier.get("title", ""))
            tier_block: Dict[str, Any] = {"title": title, "methods": []}
            for row in tier.get("rows", []):
                name = str(row["name"])
                paths = [_resolve(x) for x in row["paths"]]
                tier_block["methods"].append(_aggregate_row(name, paths, metrics))
            bundle["aggregated_dual_table"]["tiers"].append(tier_block)

    bundle["training_jsonl_inline"] = {}
    for rel in JSONL_INLINE:
        p = ROOT / rel
        if not p.is_file():
            continue
        text = p.read_text(encoding="utf-8", errors="replace")
        lines = text.splitlines()
        bundle["training_jsonl_inline"][rel] = {
            "line_count": len(lines),
            "bytes": p.stat().st_size,
            "lines": lines,
        }

    bundle["curves_csv_index"] = {}
    curves_dir = ROOT / "experiments" / "curves"
    if curves_dir.is_dir():
        for p in sorted(curves_dir.glob("*.csv")):
            bundle["curves_csv_index"][p.name] = {
                "path_relative": str(p.relative_to(ROOT)).replace("\\", "/"),
                "bytes": p.stat().st_size,
            }
    return bundle


def _md_escape_cell(s: str) -> str:
    return s.replace("|", "\\|").replace("\n", " ")


def render_markdown(bundle: Dict[str, Any]) -> str:
    """Build the same Markdown string as write_markdown (for PDF fallback)."""
    buf = StringIO()
    write_markdown_to_buffer(buf, bundle)
    return buf.getvalue()


def write_markdown_to_buffer(buf: StringIO, bundle: Dict[str, Any]) -> None:
    ts = bundle.get("generated_utc", "")
    buf.write("# 实验数据包（uav-qmix-3d）\n\n")
    buf.write(
        "本文件由 `experiments/export_experiment_data_bundle.py` 自动生成，"
        "便于归档与审阅；**非**论文正文。\n\n"
    )
    buf.write(f"- **生成时间（UTC）**：{ts}\n")
    buf.write(f"- **仓库根路径**：`{bundle.get('repo_root', '')}`\n\n")
    buf.write("---\n\n## 1. 双档五方法聚合（n=3 种子，均值 ± 样本标准差）\n\n")

    metrics: List[str] = bundle.get("aggregated_dual_table", {}).get("metrics", [])
    for tier in bundle.get("aggregated_dual_table", {}).get("tiers", []):
        title = tier.get("title", "Scenario")
        buf.write(f"### {str(title)}\n\n")
        buf.write("| 方法 | coverage_mean | success_mean | steps_mean |\n")
        buf.write("|------|---------------|--------------|------------|\n")
        for m in tier.get("methods", []):
            name = str(m.get("method", ""))
            row = [name]
            for key in metrics:
                block = m.get(key) or {}
                mu = block.get("mean")
                sd = block.get("std_sample")
                if mu is None:
                    row.append("—")
                elif key == "steps_mean":
                    row.append(f"{mu:.1f} ± {sd:.1f}")
                else:
                    row.append(f"{mu:.4f} ± {sd:.4f}")
            buf.write("| " + " | ".join(_md_escape_cell(x) for x in row) + " |\n")
        buf.write("\n<details><summary>逐方法 × 逐种子数值（展开）</summary>\n\n")
        for m in tier.get("methods", []):
            name = str(m.get("method", ""))
            buf.write(f"- **{ _md_escape_cell(name) }**\n")
            for key in metrics:
                block = m.get(key) or {}
                seeds = block.get("per_seed") or []
                files = m.get("seed_files") or []
                pairs = ", ".join(
                    f"`{Path(f).name}`:{v}"
                    for f, v in zip(files, seeds)
                    if v is not None
                )
                buf.write(f"  - **{key}**：{pairs}\n")
        buf.write("\n</details>\n\n")

    buf.write("---\n\n## 2. 每种子评估 JSON 摘要\n\n")
    buf.write(
        "| 文件 | coverage_mean | success_mean | steps_mean | 备注字段 |\n"
        "|------|-----------------|--------------|------------|----------|\n"
    )
    for fname in sorted(bundle.get("results_eval", {}).keys()):
        rec = bundle["results_eval"][fname]
        if "_error" in rec:
            buf.write(f"| `{fname}` | — | — | — | 错误：{rec['_error']} |\n")
            continue
        cov = rec.get("coverage_mean")
        suc = rec.get("success_mean")
        stp = rec.get("steps_mean")
        note = []
        if rec.get("policy"):
            note.append(f"policy={rec['policy']}")
        if rec.get("mixer_type"):
            note.append(f"mixer={rec['mixer_type']}")
        if rec.get("metrics_note"):
            note.append("含 metrics_note")
        if rec.get("metrics_source"):
            note.append(str(rec["metrics_source"]))
        buf.write(
            "| `{0}` | {1} | {2} | {3} | {4} |\n".format(
                fname,
                f"{cov:.6g}" if isinstance(cov, (int, float)) else "—",
                f"{suc:.6g}" if isinstance(suc, (int, float)) else "—",
                f"{stp:.6g}" if isinstance(stp, (int, float)) else "—",
                _md_escape_cell("; ".join(note)) if note else "—",
            )
        )

    buf.write("\n<details><summary>各 JSON 全文（折叠）</summary>\n\n")
    for fname, rec in sorted(bundle.get("results_eval", {}).items()):
        buf.write(f"#### `{fname}`\n\n```json\n")
        buf.write(json.dumps(rec, ensure_ascii=False, indent=2))
        buf.write("\n```\n\n")
    buf.write("</details>\n\n")

    buf.write("---\n\n## 3. 配置文件原文（YAML）\n\n")
    for rel, text in sorted(bundle.get("config_snapshots_text", {}).items()):
        buf.write(f"### `{rel}`\n\n```yaml\n")
        buf.write(text.rstrip() + "\n")
        buf.write("```\n\n")

    buf.write("---\n\n## 4. 训练过程 JSONL（原文逐行）\n\n")
    for rel, block in sorted(bundle.get("training_jsonl_inline", {}).items()):
        buf.write(f"### `{rel}`\n\n")
        buf.write(f"- 行数：{block.get('line_count')}，字节：{block.get('bytes')}\n\n")
        buf.write("```jsonl\n")
        buf.write("\n".join(block.get("lines") or []))
        buf.write("\n```\n\n")

    buf.write("---\n\n## 5. 学习曲线 CSV 索引\n\n")
    buf.write("| 文件名 | 相对路径 | 字节数 |\n|--------|----------|--------|\n")
    for name, meta in sorted(bundle.get("curves_csv_index", {}).items()):
        buf.write(
            f"| `{name}` | `{meta.get('path_relative','')}` | {meta.get('bytes')} |\n"
        )


def write_markdown(path: Path, bundle: Dict[str, Any]) -> None:
    buf = StringIO()
    write_markdown_to_buffer(buf, bundle)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(buf.getvalue(), encoding="utf-8")


def write_json(path: Path, bundle: Dict[str, Any]) -> None:
    out = {k: v for k, v in bundle.items() if k != "config_snapshots_text"}
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(out, f, ensure_ascii=False, indent=2)


def write_docx(path: Path, bundle: Dict[str, Any]) -> None:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.shared import Pt

    doc = Document()
    h0 = doc.add_heading("实验数据包（uav-qmix-3d）", level=0)
    h0.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph(
        "由 experiments/export_experiment_data_bundle.py 自动生成，非论文正文。"
    )
    doc.add_paragraph(f"生成时间（UTC）：{bundle.get('generated_utc', '')}")
    doc.add_paragraph(f"仓库根路径：{bundle.get('repo_root', '')}")

    metrics: List[str] = bundle.get("aggregated_dual_table", {}).get("metrics", [])
    doc.add_heading("1. 双档五方法聚合（n=3，均值±样本标准差）", level=1)
    for tier in bundle.get("aggregated_dual_table", {}).get("tiers", []):
        doc.add_heading(str(tier.get("title", "")), level=2)
        n_methods = len(tier.get("methods", []))
        tbl = doc.add_table(rows=1 + n_methods, cols=1 + len(metrics))
        hdr = tbl.rows[0].cells
        hdr[0].text = "方法"
        for j, key in enumerate(metrics, start=1):
            hdr[j].text = key
        for i, m in enumerate(tier.get("methods", []), start=1):
            row = tbl.rows[i].cells
            row[0].text = str(m.get("method", ""))
            for j, key in enumerate(metrics, start=1):
                block = m.get(key) or {}
                mu = block.get("mean")
                sd = block.get("std_sample")
                if mu is None:
                    row[j].text = "—"
                elif key == "steps_mean":
                    row[j].text = f"{mu:.1f} ± {sd:.1f}"
                else:
                    row[j].text = f"{mu:.4f} ± {sd:.4f}"
        doc.add_paragraph("逐种子明细：")
        for m in tier.get("methods", []):
            lines = [str(m.get("method", ""))]
            for key in metrics:
                block = m.get(key) or {}
                seeds = block.get("per_seed") or []
                files = m.get("seed_files") or []
                parts = ", ".join(
                    f"{Path(f).name}:{v}"
                    for f, v in zip(files, seeds)
                    if v is not None
                )
                lines.append(f"  {key}: {parts}")
            doc.add_paragraph("\n".join(lines))

    doc.add_heading("2. 每种子评估 JSON 摘要", level=1)
    n_files = len(bundle.get("results_eval", {}))
    t2 = doc.add_table(rows=1 + n_files, cols=5)
    h2 = t2.rows[0].cells
    h2[0].text = "文件"
    h2[1].text = "coverage_mean"
    h2[2].text = "success_mean"
    h2[3].text = "steps_mean"
    h2[4].text = "备注"
    for i, fname in enumerate(sorted(bundle.get("results_eval", {}).keys()), start=1):
        rec = bundle["results_eval"][fname]
        cells = t2.rows[i].cells
        cells[0].text = fname
        if "_error" in rec:
            cells[1].text = cells[2].text = cells[3].text = "—"
            cells[4].text = rec["_error"]
            continue
        cov, suc, stp = rec.get("coverage_mean"), rec.get("success_mean"), rec.get(
            "steps_mean"
        )
        cells[1].text = f"{cov:.6g}" if isinstance(cov, (int, float)) else "—"
        cells[2].text = f"{suc:.6g}" if isinstance(suc, (int, float)) else "—"
        cells[3].text = f"{stp:.6g}" if isinstance(stp, (int, float)) else "—"
        bits = []
        if rec.get("policy"):
            bits.append(f"policy={rec['policy']}")
        if rec.get("mixer_type"):
            bits.append(f"mixer={rec['mixer_type']}")
        if rec.get("metrics_source"):
            bits.append(str(rec["metrics_source"]))
        cells[4].text = "; ".join(bits) if bits else "—"

    doc.add_heading("3. 各评估 JSON 全文", level=1)
    for fname, rec in sorted(bundle.get("results_eval", {}).items()):
        doc.add_heading(fname, level=3)
        p = doc.add_paragraph()
        run = p.add_run(json.dumps(rec, ensure_ascii=False, indent=2))
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(8)

    doc.add_heading("4. 配置文件（YAML 原文）", level=1)
    for rel, text in sorted(bundle.get("config_snapshots_text", {}).items()):
        doc.add_heading(rel, level=2)
        p = doc.add_paragraph()
        r = p.add_run(text.rstrip() + "\n")
        r.font.name = "Courier New"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        r.font.size = Pt(8)

    doc.add_heading("5. 训练 JSONL 原文", level=1)
    for rel, block in sorted(bundle.get("training_jsonl_inline", {}).items()):
        doc.add_heading(rel, level=2)
        doc.add_paragraph(f"行数 {block.get('line_count')}，字节 {block.get('bytes')}")
        body = "\n".join(block.get("lines") or [])
        p = doc.add_paragraph()
        rr = p.add_run(body + ("\n" if body else ""))
        rr.font.name = "Courier New"
        rr._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        rr.font.size = Pt(8)

    doc.add_heading("6. 学习曲线 CSV 索引", level=1)
    idx = bundle.get("curves_csv_index", {})
    t3 = doc.add_table(rows=1 + len(idx), cols=3)
    z = t3.rows[0].cells
    z[0].text, z[1].text, z[2].text = "文件名", "相对路径", "字节"
    for i, (name, meta) in enumerate(sorted(idx.items()), start=1):
        c = t3.rows[i].cells
        c[0].text = name
        c[1].text = str(meta.get("path_relative", ""))
        c[2].text = str(meta.get("bytes", ""))

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def _register_pdf_cjk_font() -> str:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    candidates: List[tuple[str, int | None]] = []
    if platform.system() == "Windows":
        windir = os.environ.get("WINDIR", r"C:\Windows")
        fonts = Path(windir) / "Fonts"
        for name in (
            "msyh.ttc",
            "msyh.ttf",
            "simhei.ttf",
            "simsun.ttc",
            "msyhbd.ttf",
        ):
            p = fonts / name
            if p.is_file():
                sub: int | None = 0 if name.endswith(".ttc") else None
                candidates.append((str(p), sub))
    else:
        for p in (
            "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        ):
            if Path(p).is_file():
                candidates.append((p, 0 if p.endswith(".ttc") else None))

    for fp, sub in candidates:
        try:
            if sub is not None:
                pdfmetrics.registerFont(TTFont("CJK0", fp, subfontIndex=sub))
            else:
                pdfmetrics.registerFont(TTFont("CJK0", fp))
            return "CJK0"
        except Exception:
            continue
    return "Helvetica"


def write_pdf(path: Path, bundle: Dict[str, Any]) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        Paragraph,
        Preformatted,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    font = _register_pdf_cjk_font()
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        name="BodyCN",
        parent=styles["Normal"],
        fontName=font,
        fontSize=9,
        leading=12,
    )
    small = ParagraphStyle(
        name="SmallCN",
        parent=styles["Normal"],
        fontName=font,
        fontSize=7,
        leading=9,
    )
    h_style = ParagraphStyle(
        name="H1CN",
        parent=styles["Heading1"],
        fontName=font,
        fontSize=14,
        leading=18,
    )
    h2_style = ParagraphStyle(
        name="H2CN",
        parent=styles["Heading2"],
        fontName=font,
        fontSize=11,
        leading=14,
    )

    story: list = []
    story.append(Paragraph("实验数据包（uav-qmix-3d）", h_style))
    story.append(
        Paragraph(
            f"自动生成 · UTC {bundle.get('generated_utc','')}<br/>"
            f"仓库：{bundle.get('repo_root','')}",
            body,
        )
    )
    story.append(Spacer(1, 0.4 * cm))

    metrics: List[str] = bundle.get("aggregated_dual_table", {}).get("metrics", [])
    story.append(Paragraph("1. 双档五方法聚合（n=3）", h2_style))
    for tier in bundle.get("aggregated_dual_table", {}).get("tiers", []):
        story.append(Paragraph(str(tier.get("title", "")), body))
        data = [["方法"] + metrics]
        for m in tier.get("methods", []):
            row = [str(m.get("method", ""))]
            for key in metrics:
                block = m.get(key) or {}
                mu = block.get("mean")
                sd = block.get("std_sample")
                if mu is None:
                    row.append("—")
                elif key == "steps_mean":
                    row.append(f"{mu:.1f}±{sd:.1f}")
                else:
                    row.append(f"{mu:.4f}±{sd:.4f}")
            data.append(row)
        t = Table(data, colWidths=[2.2 * cm, 3.2 * cm, 3.2 * cm, 3.2 * cm])
        t.setStyle(
            TableStyle(
                [
                    ("FONT", (0, 0), (-1, -1), font, 8),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ]
            )
        )
        story.append(t)
        story.append(Spacer(1, 0.2 * cm))

    story.append(Paragraph("2. 每种子评估 JSON 摘要表", h2_style))
    rows = [["文件", "coverage", "success", "steps", "备注"]]
    for fname in sorted(bundle.get("results_eval", {}).keys()):
        rec = bundle["results_eval"][fname]
        if "_error" in rec:
            rows.append([fname[:24], "—", "—", "—", rec["_error"][:40]])
            continue
        cov = rec.get("coverage_mean")
        suc = rec.get("success_mean")
        stp = rec.get("steps_mean")
        note = str(rec.get("mixer_type") or rec.get("policy") or "")[:30]
        rows.append(
            [
                fname[:28],
                f"{cov:.4f}" if isinstance(cov, (int, float)) else "—",
                f"{suc:.4f}" if isinstance(suc, (int, float)) else "—",
                f"{stp:.1f}" if isinstance(stp, (int, float)) else "—",
                note,
            ]
        )
    wt = Table(
        rows,
        colWidths=[3.5 * cm, 2.2 * cm, 2.2 * cm, 2.2 * cm, 4.5 * cm],
        repeatRows=1,
    )
    wt.setStyle(
        TableStyle(
            [
                ("FONT", (0, 0), (-1, -1), font, 7),
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ]
        )
    )
    story.append(wt)
    story.append(Spacer(1, 0.3 * cm))

    story.append(
        Paragraph(
            "3. 以下为配置文件 YAML 全文（小字分块）。"
            "各评估 JSON 与 JSONL 全文见同目录 EXPERIMENT_DATA_BUNDLE.md 或 .docx。",
            small,
        )
    )
    story.append(Spacer(1, 0.2 * cm))

    for rel, text in sorted(bundle.get("config_snapshots_text", {}).items()):
        story.append(Paragraph(f"YAML：{rel}", h2_style))
        raw = text.rstrip() + "\n"
        chunk = 5000
        for i in range(0, len(raw), chunk):
            story.append(Preformatted(raw[i : i + chunk], small, maxLineLength=100))

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=1.5 * cm,
        leftMargin=1.5 * cm,
        topMargin=1.2 * cm,
        bottomMargin=1.2 * cm,
    )
    doc.build(story)


def write_pdf_matplotlib(path: Path, bundle: Dict[str, Any]) -> None:
    """Fallback PDF using matplotlib (already in requirements); plain-text pages."""
    import matplotlib.pyplot as plt
    from matplotlib.backends.backend_pdf import PdfPages

    plt.rcParams.update(
        {
            "font.sans-serif": ["Microsoft YaHei", "SimHei", "DejaVu Sans"],
            "axes.unicode_minus": False,
        }
    )
    text = render_markdown(bundle)
    lines = text.splitlines()
    lines_per_page = 42
    chunks: List[List[str]] = []
    for i in range(0, len(lines), lines_per_page):
        chunks.append(lines[i : i + lines_per_page])

    path.parent.mkdir(parents=True, exist_ok=True)
    with PdfPages(path) as pdf:
        for page_lines in chunks:
            fig, ax = plt.subplots(figsize=(8.27, 11.69))
            ax.axis("off")
            body = "\n".join(page_lines)
            ax.text(
                0.02,
                0.98,
                body,
                transform=ax.transAxes,
                fontsize=6.5,
                va="top",
                ha="left",
                family="sans-serif",
            )
            pdf.savefig(fig, bbox_inches="tight")
            plt.close(fig)


def main() -> None:
    p = argparse.ArgumentParser(description="Export experiment data bundle.")
    p.add_argument(
        "--format",
        choices=("md", "json", "docx", "pdf", "both", "all"),
        default="md",
        help="md (default), json, docx, pdf, both (=md+json), all (=md+json+docx+pdf)",
    )
    args = p.parse_args()

    bundle = collect_bundle()

    if args.format in ("md", "both", "all"):
        write_markdown(OUT_MD, bundle)
        print(f"Wrote {OUT_MD} ({OUT_MD.stat().st_size} bytes)")
    if args.format in ("json", "both", "all"):
        write_json(OUT_JSON, bundle)
        print(f"Wrote {OUT_JSON} ({OUT_JSON.stat().st_size} bytes)")
    if args.format in ("docx", "all"):
        try:
            write_docx(OUT_DOCX, bundle)
            print(f"Wrote {OUT_DOCX} ({OUT_DOCX.stat().st_size} bytes)")
        except ImportError:
            print("Skip docx: pip install python-docx")
        except Exception as exc:  # noqa: BLE001
            print(f"docx failed: {exc}")
    if args.format in ("pdf", "all"):
        try:
            write_pdf(OUT_PDF, bundle)
            print(f"Wrote {OUT_PDF} ({OUT_PDF.stat().st_size} bytes) [reportlab]")
        except ImportError:
            try:
                write_pdf_matplotlib(OUT_PDF, bundle)
                print(
                    f"Wrote {OUT_PDF} ({OUT_PDF.stat().st_size} bytes) [matplotlib fallback]"
                )
            except Exception as exc:  # noqa: BLE001
                print(f"Skip pdf: {exc}")
        except Exception as exc:  # noqa: BLE001
            try:
                write_pdf_matplotlib(OUT_PDF, bundle)
                print(
                    f"Wrote {OUT_PDF} ({OUT_PDF.stat().st_size} bytes) [matplotlib fallback after: {exc}]"
                )
            except Exception as exc2:  # noqa: BLE001
                print(f"Skip pdf: {exc}; fallback failed: {exc2}")


if __name__ == "__main__":
    main()
